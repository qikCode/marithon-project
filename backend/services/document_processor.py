"""
Document Processing Service
Handles text extraction from PDF, DOC, and DOCX files
"""

import os
import logging
from typing import Optional, Dict, Any
import PyPDF2
from docx import Document as DocxDocument
import tempfile
import subprocess
from pathlib import Path

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Service for processing and extracting text from various document formats"""
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.doc', '.docx'}
        
    def extract_text(self, file_path: str) -> str:
        """
        Extract text content from document
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format is not supported
            Exception: If extraction fails
        """
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            logger.info(f"Extracting text from {file_extension} file: {file_path}")
            
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension == '.docx':
                return self._extract_from_docx(file_path)
            elif file_extension == '.doc':
                return self._extract_from_doc(file_path)
            
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {str(e)}")
            raise
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
                        continue
            
            if not text_content:
                logger.warning("No text extracted from PDF - might be image-based")
                # Try OCR as fallback (if available)
                return self._try_ocr_extraction(file_path)
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            raise
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(' | '.join(row_text))
                
                if table_text:
                    text_content.append('\n--- Table ---\n' + '\n'.join(table_text))
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            raise
    
    def _extract_from_doc(self, file_path: str) -> str:
        """Extract text from DOC file using antiword or similar tool"""
        try:
            # Try using antiword (Linux/Mac) or textract (cross-platform)
            result = None
            
            # Method 1: Try antiword (if available)
            try:
                result = subprocess.run(
                    ['antiword', file_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout
            except (subprocess.TimeoutExpired, FileNotFoundError):
                pass
            
            # Method 2: Try catdoc (if available)
            try:
                result = subprocess.run(
                    ['catdoc', file_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout
            except (subprocess.TimeoutExpired, FileNotFoundError):
                pass
            
            # Method 3: Try python-docx2txt (fallback)
            try:
                import docx2txt
                text = docx2txt.process(file_path)
                if text and text.strip():
                    return text
            except ImportError:
                pass
            
            # Method 4: Convert to DOCX and extract (using LibreOffice if available)
            try:
                return self._convert_doc_to_docx_and_extract(file_path)
            except:
                pass
            
            raise Exception("No suitable DOC extraction method available")
            
        except Exception as e:
            logger.error(f"DOC extraction error: {str(e)}")
            raise
    
    def _convert_doc_to_docx_and_extract(self, file_path: str) -> str:
        """Convert DOC to DOCX using LibreOffice and extract text"""
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                # Convert DOC to DOCX using LibreOffice
                result = subprocess.run([
                    'libreoffice', '--headless', '--convert-to', 'docx',
                    '--outdir', temp_dir, file_path
                ], capture_output=True, timeout=60)
                
                if result.returncode != 0:
                    raise Exception("LibreOffice conversion failed")
                
                # Find the converted DOCX file
                doc_name = Path(file_path).stem
                docx_path = os.path.join(temp_dir, f"{doc_name}.docx")
                
                if os.path.exists(docx_path):
                    return self._extract_from_docx(docx_path)
                else:
                    raise Exception("Converted DOCX file not found")
                    
        except Exception as e:
            logger.error(f"DOC to DOCX conversion error: {str(e)}")
            raise
    
    def _try_ocr_extraction(self, file_path: str) -> str:
        """Try OCR extraction for image-based PDFs"""
        try:
            # This would require additional OCR libraries like pytesseract
            # For now, return a placeholder message
            logger.warning("OCR extraction not implemented - image-based PDF detected")
            return "OCR_REQUIRED: This document appears to be image-based and requires OCR processing."
            
        except Exception as e:
            logger.error(f"OCR extraction error: {str(e)}")
            return "TEXT_EXTRACTION_FAILED: Unable to extract text from this document."
    
    def get_document_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from document
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing document metadata
        """
        try:
            file_extension = Path(file_path).suffix.lower()
            metadata = {
                'file_size': os.path.getsize(file_path),
                'file_extension': file_extension,
                'pages': 0,
                'word_count': 0,
                'character_count': 0
            }
            
            if file_extension == '.pdf':
                metadata.update(self._get_pdf_metadata(file_path))
            elif file_extension == '.docx':
                metadata.update(self._get_docx_metadata(file_path))
            
            # Extract text to get word/character counts
            try:
                text = self.extract_text(file_path)
                metadata['word_count'] = len(text.split())
                metadata['character_count'] = len(text)
            except:
                pass
            
            return metadata
            
        except Exception as e:
            logger.error(f"Metadata extraction error: {str(e)}")
            return {'error': str(e)}
    
    def _get_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract PDF-specific metadata"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                metadata = {
                    'pages': len(pdf_reader.pages),
                    'title': '',
                    'author': '',
                    'subject': '',
                    'creator': ''
                }
                
                # Extract document info if available
                if pdf_reader.metadata:
                    info = pdf_reader.metadata
                    metadata.update({
                        'title': info.get('/Title', ''),
                        'author': info.get('/Author', ''),
                        'subject': info.get('/Subject', ''),
                        'creator': info.get('/Creator', ''),
                        'creation_date': info.get('/CreationDate', ''),
                        'modification_date': info.get('/ModDate', '')
                    })
                
                return metadata
                
        except Exception as e:
            logger.error(f"PDF metadata extraction error: {str(e)}")
            return {}
    
    def _get_docx_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract DOCX-specific metadata"""
        try:
            doc = DocxDocument(file_path)
            
            metadata = {
                'pages': 1,  # DOCX doesn't have fixed pages
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
            
            # Extract core properties if available
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                metadata.update({
                    'title': props.title or '',
                    'author': props.author or '',
                    'subject': props.subject or '',
                    'created': props.created.isoformat() if props.created else '',
                    'modified': props.modified.isoformat() if props.modified else ''
                })
            
            return metadata
            
        except Exception as e:
            logger.error(f"DOCX metadata extraction error: {str(e)}")
            return {}
    
    def validate_document(self, file_path: str) -> Dict[str, Any]:
        """
        Validate document and check if it's processable
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary with validation results
        """
        try:
            validation_result = {
                'is_valid': False,
                'file_exists': False,
                'is_supported_format': False,
                'is_readable': False,
                'has_text_content': False,
                'file_size': 0,
                'error_message': None
            }
            
            # Check if file exists
            if not os.path.exists(file_path):
                validation_result['error_message'] = "File does not exist"
                return validation_result
            
            validation_result['file_exists'] = True
            validation_result['file_size'] = os.path.getsize(file_path)
            
            # Check file format
            file_extension = Path(file_path).suffix.lower()
            if file_extension not in self.supported_formats:
                validation_result['error_message'] = f"Unsupported format: {file_extension}"
                return validation_result
            
            validation_result['is_supported_format'] = True
            
            # Try to read the file
            try:
                with open(file_path, 'rb') as f:
                    f.read(1024)  # Try to read first 1KB
                validation_result['is_readable'] = True
            except Exception as e:
                validation_result['error_message'] = f"File is not readable: {str(e)}"
                return validation_result
            
            # Try to extract some text
            try:
                text = self.extract_text(file_path)
                if text and len(text.strip()) > 10:  # At least 10 characters
                    validation_result['has_text_content'] = True
                    validation_result['is_valid'] = True
                else:
                    validation_result['error_message'] = "No readable text content found"
            except Exception as e:
                validation_result['error_message'] = f"Text extraction failed: {str(e)}"
            
            return validation_result
            
        except Exception as e:
            logger.error(f"Document validation error: {str(e)}")
            return {
                'is_valid': False,
                'error_message': f"Validation failed: {str(e)}"
            }
